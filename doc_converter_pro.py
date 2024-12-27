from pdf2docx import Converter
import pytesseract
from pdf2image import convert_from_path
from docx import Document
from PIL import Image
import speech_recognition as sr
import os
from concurrent.futures import ThreadPoolExecutor

# Configure Tesseract path for Windows
if os.name == 'nt':
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Function to convert PDF to Word preserving original formatting
def pdf_to_word(pdf_path, word_path):
    cv = Converter(pdf_path)
    cv.convert(word_path, start=0, end=None)
    cv.close()

# Function to convert JPG to Word with OCR capabilities
def jpg_to_word_with_ocr(jpg_path, word_path):
    image = Image.open(jpg_path)
    text = pytesseract.image_to_string(image)
    # Remove any XML incompatible characters
    text = ''.join(c for c in text if c.isprintable())
    doc = Document()
    doc.add_paragraph(text)
    doc.save(word_path)

# Function to convert audio to text
def audio_to_text(audio_path, text_path):
    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_path) as source:
        audio_data = recognizer.record(source)
        text = recognizer.recognize_google(audio_data)
    with open(text_path, 'w') as file:
        file.write(text)

# Function to convert PDF to Word with OCR capabilities using multithreading
def pdf_to_word_with_ocr_multithreaded(pdf_path, word_path):
    images = convert_from_path(pdf_path)
    doc = Document()

    def process_image(image):
        text = pytesseract.image_to_string(image)
        # Remove any XML incompatible characters
        text = ''.join(c for c in text if c.isprintable())
        return text

    with ThreadPoolExecutor() as executor:
        results = executor.map(process_image, images)

    for text in results:
        doc.add_paragraph(text)

    doc.save(word_path)

# Main function to prompt user and perform the desired action
def main():
    while True:
        print("Choose an option:")
        print("1. Convert PDF to Word")
        print("2. Convert JPG to Word")
        print("3. Convert Audio to Text")
        print("0. Exit")
        choice = input("Enter your choice: ")

        if choice == '0':
            break
        elif choice in ['1', '2', '3']:
            file_path = input("Enter the file path: ")
            file_dir, file_name = os.path.split(file_path)
            file_base, file_ext = os.path.splitext(file_name)

            if choice == '1':
                output_path = os.path.join(file_dir, f"{file_base}.docx")
                pdf_to_word_with_ocr_multithreaded(file_path, output_path)
                print(f"PDF has been successfully converted to Word and saved as {output_path}.")
            elif choice == '2':
                output_path = os.path.join(file_dir, f"{file_base}.docx")
                jpg_to_word_with_ocr(file_path, output_path)
                print(f"JPG has been successfully converted to Word and saved as {output_path}.")
            elif choice == '3':
                output_path = os.path.join(file_dir, f"{file_base}.txt")
                audio_to_text(file_path, output_path)
                print(f"Audio has been successfully converted to text and saved as {output_path}.")
        else:
            print("Invalid choice. Please try again.")

if __name__ == "__main__":
    main()
